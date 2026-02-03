"""
문서 로딩 및 파싱 모듈

이 모듈은 다양한 형식의 문서(PDF, DOCX, Markdown, TXT)를 로드하고
텍스트를 추출하는 기능을 제공합니다.
"""

import os
from pathlib import Path
from typing import List, Optional
from dataclasses import dataclass, field
from datetime import datetime

from ..utils.logger import get_logger


logger = get_logger(__name__)


@dataclass
class DocumentMetadata:
    """
    문서 메타데이터 클래스
    
    Attributes:
        filename: 파일명
        file_type: 파일 형식 (pdf, docx, md, txt)
        file_path: 파일 전체 경로
        file_size: 파일 크기 (bytes)
        created_at: 생성 시간
        modified_at: 수정 시간
        page_count: 페이지 수 (PDF의 경우)
    """
    filename: str
    file_type: str
    file_path: str
    file_size: int
    created_at: Optional[datetime] = None
    modified_at: Optional[datetime] = None
    page_count: Optional[int] = None
    extra: dict = field(default_factory=dict)


@dataclass
class Document:
    """
    로드된 문서 클래스
    
    Attributes:
        content: 문서 텍스트 내용
        metadata: 문서 메타데이터
    """
    content: str
    metadata: DocumentMetadata


class DocumentLoader:
    """
    문서 로더 클래스
    
    다양한 형식의 문서를 로드하고 텍스트를 추출합니다.
    """
    
    SUPPORTED_EXTENSIONS = {'.pdf', '.docx', '.md', '.txt'}
    
    def __init__(self):
        """DocumentLoader 초기화"""
        self._parsers = {
            '.pdf': self._parse_pdf,
            '.docx': self._parse_docx,
            '.md': self._parse_markdown,
            '.txt': self._parse_txt,
        }
    
    def load_file(self, file_path: str) -> Document:
        """
        단일 파일을 로드합니다.
        
        Args:
            file_path: 파일 경로
            
        Returns:
            Document: 로드된 문서
            
        Raises:
            FileNotFoundError: 파일이 존재하지 않는 경우
            ValueError: 지원하지 않는 파일 형식인 경우
        """
        path = Path(file_path)
        
        if not path.exists():
            raise FileNotFoundError(f"파일을 찾을 수 없습니다: {file_path}")
        
        ext = path.suffix.lower()
        if ext not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(f"지원하지 않는 파일 형식입니다: {ext}")
        
        # 메타데이터 생성
        stat = path.stat()
        metadata = DocumentMetadata(
            filename=path.name,
            file_type=ext[1:],  # 점(.) 제거
            file_path=str(path.absolute()),
            file_size=stat.st_size,
            created_at=datetime.fromtimestamp(stat.st_ctime),
            modified_at=datetime.fromtimestamp(stat.st_mtime),
        )
        
        # 파서 호출
        parser = self._parsers[ext]
        content, extra_metadata = parser(path)
        
        # 추가 메타데이터 병합
        if extra_metadata:
            metadata.extra.update(extra_metadata)
            if 'page_count' in extra_metadata:
                metadata.page_count = extra_metadata['page_count']
        
        logger.info(f"문서 로드 완료: {path.name}, 크기: {len(content)} 문자")
        
        return Document(content=content, metadata=metadata)
    
    def load_directory(
        self,
        directory_path: str,
        recursive: bool = True
    ) -> List[Document]:
        """
        디렉토리의 모든 문서를 로드합니다.
        
        Args:
            directory_path: 디렉토리 경로
            recursive: 하위 디렉토리 포함 여부
            
        Returns:
            List[Document]: 로드된 문서 목록
        """
        path = Path(directory_path)
        
        if not path.exists():
            raise FileNotFoundError(f"디렉토리를 찾을 수 없습니다: {directory_path}")
        
        if not path.is_dir():
            raise ValueError(f"디렉토리가 아닙니다: {directory_path}")
        
        documents = []
        pattern = '**/*' if recursive else '*'
        
        for file_path in path.glob(pattern):
            if file_path.is_file() and file_path.suffix.lower() in self.SUPPORTED_EXTENSIONS:
                try:
                    doc = self.load_file(str(file_path))
                    documents.append(doc)
                except Exception as e:
                    logger.error(f"문서 로드 실패: {file_path}, 오류: {e}")
        
        logger.info(f"디렉토리 스캔 완료: {len(documents)}개 문서 로드")
        
        return documents
    
    def _parse_pdf(self, path: Path) -> tuple[str, dict]:
        """
        PDF 파일을 파싱합니다.
        
        Args:
            path: 파일 경로
            
        Returns:
            tuple: (텍스트 내용, 추가 메타데이터)
        """
        try:
            from pypdf import PdfReader
            
            reader = PdfReader(str(path))
            pages = []
            
            for i, page in enumerate(reader.pages):
                text = page.extract_text()
                if text:
                    pages.append(f"[페이지 {i + 1}]\n{text}")
            
            content = "\n\n".join(pages)
            metadata = {
                'page_count': len(reader.pages),
                'pdf_info': reader.metadata._data if reader.metadata else {}
            }
            
            return content, metadata
            
        except ImportError:
            logger.warning("pypdf 패키지가 설치되지 않았습니다. PDF 파싱을 건너뜁니다.")
            return "", {'error': 'pypdf not installed'}
        except Exception as e:
            logger.error(f"PDF 파싱 오류: {path}, {e}")
            return "", {'error': str(e)}
    
    def _parse_docx(self, path: Path) -> tuple[str, dict]:
        """
        DOCX 파일을 파싱합니다.
        
        Args:
            path: 파일 경로
            
        Returns:
            tuple: (텍스트 내용, 추가 메타데이터)
        """
        try:
            from docx import Document as DocxDocument
            
            doc = DocxDocument(str(path))
            paragraphs = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            
            # 테이블 내용도 추출
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        paragraphs.append(row_text)
            
            content = "\n\n".join(paragraphs)
            metadata = {
                'paragraph_count': len(doc.paragraphs),
                'table_count': len(doc.tables),
            }
            
            return content, metadata
            
        except ImportError:
            logger.warning("python-docx 패키지가 설치되지 않았습니다. DOCX 파싱을 건너뜁니다.")
            return "", {'error': 'python-docx not installed'}
        except Exception as e:
            logger.error(f"DOCX 파싱 오류: {path}, {e}")
            return "", {'error': str(e)}
    
    def _parse_markdown(self, path: Path) -> tuple[str, dict]:
        """
        Markdown 파일을 파싱합니다.
        
        Args:
            path: 파일 경로
            
        Returns:
            tuple: (텍스트 내용, 추가 메타데이터)
        """
        try:
            content = path.read_text(encoding='utf-8')
            
            # 헤더 수 계산
            headers = [line for line in content.split('\n') if line.startswith('#')]
            
            metadata = {
                'header_count': len(headers),
                'line_count': len(content.split('\n')),
            }
            
            return content, metadata
            
        except Exception as e:
            logger.error(f"Markdown 파싱 오류: {path}, {e}")
            return "", {'error': str(e)}
    
    def _parse_txt(self, path: Path) -> tuple[str, dict]:
        """
        TXT 파일을 파싱합니다.
        
        Args:
            path: 파일 경로
            
        Returns:
            tuple: (텍스트 내용, 추가 메타데이터)
        """
        try:
            # 여러 인코딩 시도
            encodings = ['utf-8', 'cp949', 'euc-kr', 'latin-1']
            content = None
            used_encoding = None
            
            for encoding in encodings:
                try:
                    content = path.read_text(encoding=encoding)
                    used_encoding = encoding
                    break
                except UnicodeDecodeError:
                    continue
            
            if content is None:
                raise ValueError("지원되는 인코딩으로 파일을 읽을 수 없습니다.")
            
            metadata = {
                'encoding': used_encoding,
                'line_count': len(content.split('\n')),
            }
            
            return content, metadata
            
        except Exception as e:
            logger.error(f"TXT 파싱 오류: {path}, {e}")
            return "", {'error': str(e)}


# 편의 함수
def load_document(file_path: str) -> Document:
    """
    단일 문서를 로드하는 편의 함수
    
    Args:
        file_path: 파일 경로
        
    Returns:
        Document: 로드된 문서
    """
    loader = DocumentLoader()
    return loader.load_file(file_path)


def load_documents_from_directory(
    directory_path: str,
    recursive: bool = True
) -> List[Document]:
    """
    디렉토리의 문서들을 로드하는 편의 함수
    
    Args:
        directory_path: 디렉토리 경로
        recursive: 하위 디렉토리 포함 여부
        
    Returns:
        List[Document]: 로드된 문서 목록
    """
    loader = DocumentLoader()
    return loader.load_directory(directory_path, recursive)
